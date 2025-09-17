import streamlit as st
import pandas as pd
import os
from docx import Document
from io import BytesIO
import zipfile

# Function to replace placeholders in Word
def replace_placeholders(template, row):
    # Handle paragraphs
    for p in template.paragraphs:
        text = p.text
        for key, val in row.items():
            placeholder = f"{{{key}}}"
            text = text.replace(placeholder, str(val))
        if text != p.text:
            for r in p.runs:
                r.text = ""
            p.add_run(text)

    # Handle tables
    for table in template.tables:
        for row_cells in table.rows:
            for cell in row_cells.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for key, val in row.items():
                        placeholder = f"{{{key}}}"
                        text = text.replace(placeholder, str(val))
                    if text != p.text:
                        for r in p.runs:
                            r.text = ""
                        p.add_run(text)
    return template

st.title("📄 Experience Letter Generator")

# Upload Word template
template_file = st.file_uploader("Upload Word Template (.docx)", type=["docx"])

# Upload Excel file
excel_file = st.file_uploader("Upload Excel File (.xlsx)", type=["xlsx"])

if template_file and excel_file:
    df = pd.read_excel(excel_file)
    df.columns = df.columns.str.strip().str.lower()  # Normalize headers
    st.write("### Detected Columns:", df.columns.tolist())
    st.dataframe(df.head())

    if st.button("Generate Letters"):
        template_bytes = template_file.read()
        count = 0
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "w") as zipf:
            for idx, row in df.iterrows():
                doc = Document(BytesIO(template_bytes))
                data = {
                    "name": row.get("name", ""),
                    "empcode": row.get("empcode", ""),
                    "designation": row.get("designation", ""),
                    "department": row.get("department", ""),
                    "doj": row.get("doj").strftime("%d-%m-%Y") if pd.notna(row.get("doj")) else "",
                    "lwd": row.get("lwd").strftime("%d-%m-%Y") if pd.notna(row.get("lwd")) else ""
                }

                # Replace placeholders
                doc = replace_placeholders(doc, data)

                # Save each doc to memory
                file_stream = BytesIO()
                filename = f"{data['empcode']}_{data['name']}.docx".replace(" ", "_")
                doc.save(file_stream)
                file_stream.seek(0)
                zipf.writestr(filename, file_stream.read())
                count += 1

        zip_buffer.seek(0)

        st.success(f"✅ {count} letters generated.")
        st.download_button(
            label="⬇️ Download All Letters (ZIP)",
            data=zip_buffer,
            file_name="Generated_Letters.zip",
            mime="application/zip"
        )
